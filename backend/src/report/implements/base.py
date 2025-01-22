from io import BytesIO

from docx import Document

from src.report.core import ReportGenerator


class BaseReportGenerator(ReportGenerator):
    def generate(self, image: BytesIO) -> Document:
        document = Document()
        document.add_heading(text="Анализ рекламной компании", level=0)

        document.add_paragraph(
            'Базовая статистика', style='List Number'
        )
        document.add_picture(image_path_or_stream=image)
        return document
